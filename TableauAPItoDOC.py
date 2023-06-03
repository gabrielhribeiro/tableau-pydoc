#maintanance
#I ve anonymize some names as dataframe manipulations, but this snippet its fully functional, Enjoy


import tableauserverclient as TSC
import pandas as pd
import io
from io import BytesIO
from docx import Document
from docx.shared import Pt
from datetime import datetime

server_url = 'https://tableau.domain.com'
tableau_auth = TSC.PersonalAccessTokenAuth(token_name='reports', personal_access_token='')
server = TSC.Server(server_url, use_server_version=True)
       
#function to get ID by name of dash
def get_workbook_id(desired_workbook_name):
    with server.auth.sign_in(tableau_auth):
        workbooks, pagination_item = server.workbooks.get()
        desired_workbook = next((workbook for workbook in workbooks if workbook.name == desired_workbook_name), None)
        if desired_workbook:
            return desired_workbook.id
    return None

def get_view_id(server, tableau_auth, desired_workbook_name, desired_view_name):
    workbook_id = get_workbook_id(desired_workbook_name)
    if workbook_id:
        with server.auth.sign_in(tableau_auth):
            workbook = server.workbooks.get_by_id(workbook_id)
            server.workbooks.populate_views(workbook)
            desired_view = next((view for view in workbook.views if view.name == desired_view_name), None)
            if desired_view:
                return desired_view.id
    return None

def getvID(variable):
    return get_view_id(server, tableau_auth, 'Weekly Report Helper - GAP', variable)

#Get Variables
print("REPORT AUTOMATION") 
print("Do your analysis using the dashboard https://tableau.domain.com/views/report/")
print("To proceed with the outpout, I have some questions:")
print("Please, give the inputs for range of dates as example: 2023-05-22 12:00:00 AM")
StartDate = input("Start Date:")
EndDate = input("Final Date:")
Description = input("Paste the target for what you would like to report:")

#get id and save it to a variable to be used inside at the connections
idHelper = getvID('Report - API')
imageHelper = getvID('Report Image - Dash')
MonitoredBlockBrand = getvID('Monitored')
BDA = getvID('Target1')
BVN = getvID('Target1')
BVH = getvID('Target1')
BVC = getvID('Target1')

#connection
with server.auth.sign_in(tableau_auth):
    
   
    MonitorBDA = server.views.get_by_id(BDA)
    csv_MonitorBDA = TSC.CSVRequestOptions(maxage=5)
    csv_MonitorBDA.vf('StartDate', StartDate)
    csv_MonitorBDA.vf('EndDate', EndDate)
     # Populate the view with CSV data
    server.views.populate_csv(MonitorBDA, csv_MonitorBDA)
    
    # Retrieve the CSV data for the view
    csvs_MonitorBDA = b''.join(MonitorBDA.csv)
    dfMonitorBDA = pd.read_csv(io.StringIO(csvs_MonitorBDA.decode('utf-8')))
    
    MonitorBVN = server.views.get_by_id(BVN)
    csv_MonitorBVN = TSC.CSVRequestOptions(maxage=5)
    csv_MonitorBVN.vf('StartDate', StartDate)
    csv_MonitorBVN.vf('EndDate', EndDate)
     # Populate the view with CSV data
    server.views.populate_csv(MonitorBVN, csv_MonitorBVN)
    
    # Retrieve the CSV data for the view
    csvs_MonitorBVN = b''.join(MonitorBVN.csv)
    dfMonitorBVN = pd.read_csv(io.StringIO(csvs_MonitorBVN.decode('utf-8')))
    
    MonitorBVH = server.views.get_by_id(BVH)
    csv_MonitorBVH = TSC.CSVRequestOptions(maxage=5)
    csv_MonitorBVH.vf('StartDate', StartDate)
    csv_MonitorBVH.vf('EndDate', EndDate)
     # Populate the view with CSV data
    server.views.populate_csv(MonitorBVH, csv_MonitorBVH)
    
    # Retrieve the CSV data for the view
    csvs_MonitorBVH = b''.join(MonitorBVH.csv)
    dfMonitorBVH = pd.read_csv(io.StringIO(csvs_MonitorBVH.decode('utf-8')))
    
    MonitorBVC = server.views.get_by_id(BVC)
    csv_MonitorBVC = TSC.CSVRequestOptions(maxage=5)
    csv_MonitorBVC.vf('StartDate', StartDate)
    csv_MonitorBVC.vf('EndDate', EndDate)
    # Populate the view with CSV data
    server.views.populate_csv(MonitorBVC, csv_MonitorBVC)
    
    # Retrieve the CSV data for the view
    csvs_MonitorBVC = b''.join(MonitorBVC.csv)
    dfMonitorBVC = pd.read_csv(io.StringIO(csvs_MonitorBVC.decode('utf-8')))

    MonitorBlocked = server.views.get_by_id(MonitoredBlockBrand)
    csv_MBlocked = TSC.CSVRequestOptions(maxage=5)
    csv_MBlocked.vf('StartDate', StartDate)
    csv_MBlocked.vf('EndDate', EndDate)

    server.views.populate_csv(MonitorBlocked, csv_MBlocked)
    
    # Retrieve the CSV data for the view
    csv_MonitorBlocked = b''.join(MonitorBlocked.csv)
    dfMonitorBlocked = pd.read_csv(io.StringIO(csv_MonitorBlocked.decode('utf-8')))
    valMonitorBlocked = dfMonitorBlocked['column1'][0]
    
    view_item_helper = server.views.get_by_id(idHelper)
    # Set the CSV request options
    csv_req_option = TSC.CSVRequestOptions(maxage=5)
    csv_req_option.vf('StartDate', StartDate)
    csv_req_option.vf('EndDate', EndDate)
    csv_req_option.vf('DescriptionSelector', Description)
    csv_req_option.vf('DescSelector', 'Include')

    # Populate the view with CSV data
    server.views.populate_csv(view_item_helper, csv_req_option)
    
    # Retrieve the CSV data for the view
    csv_data = b''.join(view_item_helper.csv)
    
    dfReApi = pd.read_csv(io.StringIO(csv_data.decode('utf-8')))
    #print("dataframe", dfReApi) #For maintainance uncoment this line

####Loaded data to be inserted into the .docx####     
    # Create a DataFrame from the CSV data

    #Description are using the value inputed -> Description - {{description}}.  
 
    #Total Events - {{Affectedsessions}} 
    last_valueReApi = dfReApi['Measure Values'].iloc[-1]
    CategoryReApi = dfReApi['Category'].iloc[2]
    FormReApi = dfReApi['column2'].iloc[2]
    DestReApi = dfReApi['column2'].iloc[2]
    VendorReApi = dfReApi['column24'].iloc[2]
    SensitiveDataReApi = dfReApi['column23'].iloc[2]
    
    #Table
    Table_affected = dfReApi[['column21', 'column22']]
    table_section = Table_affected.drop(Table_affected.index[-1])
    table_section
    
    ##IMAGEM##
    view_image_helper = server.views.get_by_id(imageHelper)
    
    # set the image request option
    image_req_option = TSC.ImageRequestOptions(imageresolution=TSC.ImageRequestOptions.Resolution.High, maxage=1)

    # (optional) set a view filter    
    image_req_option.vf('StartDate', StartDate)
    image_req_option.vf('EndDate', EndDate)
    image_req_option.vf('DescriptionSelector', Description)
    image_req_option.vf('DescSelector', 'Include')

    # retrieve the image for the view
    server.views.populate_image(view_image_helper, image_req_option)
    image_stream = BytesIO(view_image_helper.image)
    # Populate and save the view image as 'view_image.png'
    with open('./Desktop/view_image.png', 'wb') as f:
        f.write(view_image_helper.image)

current_date = datetime.now().strftime("%Y-%m-%d")
document = Document()

def substitute_values_in_docx(template_path, output_path, substitutions, substitutions_img, table_section, dfMonitorBDA, dfMonitorBVN, dfMonitorBVH, dfMonitorBVC):
    # Load the template document
    doc = Document(template_path)

        # Perform the substitution of placeholders
    for paragraph in doc.paragraphs:
        for key, value in substitutions.items():
            if key in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if key in inline[i].text:
                        text = inline[i].text.replace(key, value)
                        inline[i].text = text

    # Iterate through all tables in the document
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Check if the cell contains any placeholders to be replaced
                if any(sub in cell.text for sub in substitutions):
                    # Replace the placeholders with the corresponding values
                    for sub, value in substitutions.items():
                        cell.text = cell.text.replace(sub, value)

    # Insert the image if the placeholder exists
    if '[IMAGECOUNT]' in substitutions_img:
        image_placeholder = '[IMAGECOUNT]'
        image_path = substitutions_img[image_placeholder]
        # Find the paragraph containing the image placeholder
        for paragraph in doc.paragraphs:
            if image_placeholder in paragraph.text:
                # Remove the existing run containing the placeholder
                for run in paragraph.runs:
                    if image_placeholder in run.text:
                        paragraph._p.remove(run._r)
                # Add the new run with the image
                run = paragraph.add_run()
                run.add_picture(image_path, width=Pt(500), height=Pt(300))

                
    table = doc.tables[0]  # Assuming the table is the first table in the document
    print(table)
    
    # Iterating over the data and filling the table
    for i, row in table_section.iterrows():
        cells = table.add_row().cells
        cells[0].text = row['Column1']
        cells[1].text = row['Column12'] 
    
    table = doc.tables[1]  # Assuming the table is the first table in the document
    print(table)
    
    # Iterating over the data and filling the table
    for i, row in dfMonitorBDA.iterrows():
        cells = table.add_row().cells
        cells[0].text = row['Column1']
        cells[1].text = row['Column12']
    
    table = doc.tables[2]  # Assuming the table is the first table in the document
    print(table)
    
    # Iterating over the data and filling the table
    for i, row in dfMonitorBVN.iterrows():
        cells = table.add_row().cells
        cells[0].text = row['Column12']
        cells[1].text = str(row['Column11'])
            
    table = doc.tables[3]  # Assuming the table is the first table in the document
    print(table)
    
    # Iterating over the data and filling the table
    for i, row in dfMonitorBVH.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row['VColumn112'])
        cells[1].text = str(row['Column12'])
                    
    table = doc.tables[4]  # Assuming the table is the first table in the document
    print(table)
    
    # Iterating over the data and filling the table
    for i, row in dfMonitorBVC.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row['Column12'])
        cells[1].text = str(row['Column12'])
                            
    # Save the modified document
    doc.save(output_path)

# Example usage
template = "./Desktop/my_word_template.docx"
output = "./Desktop/Report_"+current_date+".docx"

substitutions = {
    '[DESCRIPTION]': Description,
    '[VALUEAF]': str(last_valueReApi),
    '[VENDNAME]': str(VendorReApi),
    '[CATEGORY]': CategoryReApi,
    '[NETURI]': str(DestReApi),
    '[SENSIDATA]': SensitiveDataReApi,
    '[FORMDATA]': FormReApi,
    '[NSESSIONS]': valMonitorBlocked
}

substitutions_img = {
    '[IMAGECOUNT]': image_stream 
}


try:
    substitute_values_in_docx(template, output, substitutions, substitutions_img, table_section, dfMonitorBDA, dfMonitorBVN, dfMonitorBVH, dfMonitorBVC)
    print("Execution completed successfully. You can find the report at:",output)
except Exception as e:
    print("An error occurred during execution:", str(e))
